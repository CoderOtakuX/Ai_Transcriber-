from fastapi import FastAPI, File, UploadFile, BackgroundTasks
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from faster_whisper import WhisperModel
from docx import Document
import os
import uuid
import time

app = FastAPI()

# Use a much faster model
model = WhisperModel("small", device="cpu", compute_type="int8")

UPLOAD_DIR = "uploads"
RESULTS_DIR = "results"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(RESULTS_DIR, exist_ok=True)

# Shared state for progress and live transcript
transcription_progress = {}

def transcribe_and_save(audio_path, docx_path, file_id):
    segments_gen, info = model.transcribe(audio_path)
    doc = Document()
    doc.add_heading("Transcription", 0)
    doc.add_paragraph(f"Detected language: {info.language} (probability: {info.language_probability:.2f})")
    doc.add_paragraph("")
    transcript = []
    transcript_no_timestamps = []
    start_time = time.time()
    segment_count = 0
    import soundfile as sf
    f = sf.SoundFile(audio_path)
    duration = len(f) / f.samplerate
    est_total_segments = max(1, int(duration // 3))
    for segment in segments_gen:
        segment_count += 1
        transcript.append(f"[{segment.start:.2f}s -> {segment.end:.2f}s] {segment.text}")
        transcript_no_timestamps.append(segment.text)
        elapsed = time.time() - start_time
        avg_time_per_segment = elapsed / segment_count
        est_time_left = max(0, avg_time_per_segment * (est_total_segments - segment_count))
        transcription_progress[file_id] = {
            "progress": min(100, int((segment_count / est_total_segments) * 100)),
            "partial_transcript": "\n".join(transcript),
            "estimated_time_left": int(est_time_left),
            "done": False
        }
    for line in transcript:
        doc.add_paragraph(line)
    # Remove summary generation
    doc.save(docx_path)
    # Save no-timestamp version
    doc2 = Document()
    doc2.add_heading("Transcription", 0)
    doc2.add_paragraph(f"Detected language: {info.language} (probability: {info.language_probability:.2f})")
    doc2.add_paragraph("")
    for line in transcript_no_timestamps:
        doc2.add_paragraph(line)
    doc2.save(docx_path.replace('.docx', '_notimestamps.docx'))
    transcription_progress[file_id] = {
        "progress": 100,
        "partial_transcript": "\n".join(transcript),
        "estimated_time_left": 0,
        "done": True
    }

@app.post("/transcribe/")
async def transcribe_audio(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    audio_path = os.path.join(UPLOAD_DIR, f"{file_id}_{file.filename}")
    with open(audio_path, "wb") as f:
        f.write(await file.read())
    docx_path = os.path.join(RESULTS_DIR, f"{file_id}.docx")
    transcription_progress[file_id] = {"progress": 0, "partial_transcript": "", "estimated_time_left": 0, "done": False}
    background_tasks.add_task(transcribe_and_save, audio_path, docx_path, file_id)
    return {"file_id": file_id}

@app.get("/progress/{file_id}")
def get_progress(file_id: str):
    prog = transcription_progress.get(file_id)
    if not prog:
        return JSONResponse({"error": "No such transcription."}, status_code=404)
    return prog

@app.get("/download/{file_id}")
def download_docx(file_id: str, notimestamps: bool = False):
    if notimestamps:
        docx_path = os.path.join(RESULTS_DIR, f"{file_id}_notimestamps.docx")
        filename = "transcription_no_timestamps.docx"
    else:
        docx_path = os.path.join(RESULTS_DIR, f"{file_id}.docx")
        filename = "transcription.docx"
    if not os.path.exists(docx_path):
        return {"error": "Transcription not ready. Please try again later."}
    return FileResponse(docx_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=filename)

@app.get("/")
def main():
    return HTMLResponse(
        """
        <html>
            <head>
                <title>AI Audio Transcription</title>
                <style>
                body { background: #181a1b; color: #eee; font-family: Arial, sans-serif; }
                .light-mode { background: #fafafa; color: #222; }
                #progress-bar { width: 100%; background: #333; height: 24px; border-radius: 8px; margin-top: 10px; }
                #progress-bar-inner { height: 100%; width: 0%; background: #4caf50; border-radius: 8px; transition: width 0.5s; color: #fff; text-align: center; }
                #live-transcript { border: 1px solid #444; padding: 10px; margin-top: 10px; height: 200px; overflow-y: auto; background: #222; white-space: pre-line; }
                .light-mode #progress-bar { background: #eee; }
                .light-mode #progress-bar-inner { background: #4caf50; color: #222; }
                .light-mode #live-transcript { background: #fafafa; border: 1px solid #ccc; color: #222; }
                .toggle-btn { margin: 10px 0; padding: 6px 16px; border-radius: 6px; border: none; background: #444; color: #fff; cursor: pointer; }
                .light-mode .toggle-btn { background: #ddd; color: #222; }
                </style>
            </head>
            <body>
                <button class="toggle-btn" onclick="toggleMode()">Toggle Dark/Light Mode</button>
                <h1>Upload Audio File for Transcription</h1>
                <form id="upload-form" enctype="multipart/form-data">
                    <input type="file" name="file" accept="audio/*" required />
                    <button type="submit">Upload & Transcribe</button>
                </form>
                <div id="status"></div>
                <div id="progress-bar"><div id="progress-bar-inner"></div></div>
                <div>Estimated time left: <span id="eta">-</span> seconds</div>
                <div id="download-link"></div>
                <div id="live-transcript"></div>
                <script>
                function toggleMode() {
                    document.body.classList.toggle('light-mode');
                }
                const form = document.getElementById('upload-form');
                let pollInterval = null;
                form.onsubmit = async (e) => {
                    e.preventDefault();
                    const formData = new FormData(form);
                    document.getElementById('status').innerText = 'Uploading and transcribing...';
                    document.getElementById('download-link').innerHTML = '';
                    document.getElementById('live-transcript').innerText = '';
                    document.getElementById('progress-bar-inner').style.width = '0%';
                    document.getElementById('progress-bar-inner').innerText = '';
                    document.getElementById('eta').innerText = '-';
                    const resp = await fetch('/transcribe/', { method: 'POST', body: formData });
                    const data = await resp.json();
                    if (data.file_id) {
                        document.getElementById('status').innerText = 'Transcription in progress. Please wait...';
                        pollInterval = setInterval(async () => {
                            const res = await fetch(`/progress/${data.file_id}`);
                            if (res.ok) {
                                const prog = await res.json();
                                document.getElementById('progress-bar-inner').style.width = prog.progress + '%';
                                document.getElementById('progress-bar-inner').innerText = prog.progress + '%';
                                document.getElementById('eta').innerText = prog.estimated_time_left;
                                document.getElementById('live-transcript').innerText = prog.partial_transcript;
                                if (prog.done) {
                                    document.getElementById('status').innerText = 'Transcription complete!';
                                    document.getElementById('download-link').innerHTML = `
                                        <a href="/download/${data.file_id}">Download Word File (with timestamps)</a> |
                                        <a href="/download/${data.file_id}?notimestamps=true">Download Word File (text only)</a>
                                    `;
                                    clearInterval(pollInterval);
                                }
                            }
                        }, 1000);
                    } else {
                        document.getElementById('status').innerText = 'Error uploading file.';
                    }
                };
                </script>
            </body>
        </html>
        """
    ) 